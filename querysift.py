import argparse
import os
import json
import pandas as pd
import re
import string
import torch
import numpy as np
from tqdm import tqdm
from transformers import AutoModel, AutoTokenizer, pipeline
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Setup command line arguments
parser = argparse.ArgumentParser(description="Process and analyze text data using NLP models.")
parser.add_argument("--chunk_size", type=int, default=1000)
parser.add_argument("--chunk_overlap", type=int, default=500)
parser.add_argument("--highlight_limit", type=int, default=5)
parser.add_argument("--top_results_limit", type=int, default=50)
parser.add_argument("--minimum_merge_overlap", type=int, default=10)
parser.add_argument("--embedding_model", type=str, default="jinaai/jina-embeddings-v3")
parser.add_argument("--qa_model", type=str, default="deepset/gelectra-large-germanquad")
args = parser.parse_args()

WINDOW_SIZE = args.chunk_size
OVERLAP = args.chunk_overlap
HIGHLIGHT_LIMIT = args.highlight_limit
TOP_EMB_LIMIT = args.top_results_limit
MIN_OVERLAP = args.minimum_merge_overlap
model_name = args.embedding_model
qa_model_name = args.qa_model

# Load models
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModel.from_pretrained(model_name, trust_remote_code=True)
nlp = pipeline('question-answering', model=qa_model_name, tokenizer=qa_model_name, topk=HIGHLIGHT_LIMIT)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)

# Load category questions
with open('category_questions.json', 'r', encoding='utf-8') as f:
    category_questions = json.load(f)

def extract_questions_with_categories(category_dict, path=None):
    if path is None:
        path = []
    queries = []
    for key, value in category_dict.items():
        if isinstance(value, dict):
            queries.extend(extract_questions_with_categories(value, path + [key]))
        elif isinstance(value, list):
            for query in value:
                queries.append((query, path + [key]))
    return queries

query_data = extract_questions_with_categories(category_questions)
queries, paths = zip(*query_data)
max_depth = max(len(p) for p in paths)
paths = [p + [None]*(max_depth - len(p)) for p in paths]
category_columns = [f"Category_Level_{i+1}" for i in range(max_depth)]
query_df = pd.DataFrame(paths, columns=category_columns)
query_df.insert(0, "Query", queries)

# Encode texts
def encode_texts(texts, batch_size=16):
    embeddings = []
    with torch.no_grad():
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i+batch_size]
            batch_outputs = model.encode(batch, task="text-matching")
            embeddings.extend(batch_outputs)
    return torch.tensor(embeddings).to(device)

def calculate_similarity(query_embeddings, passage_embeddings):
    return cosine_similarity(query_embeddings.cpu().numpy(), passage_embeddings.cpu().numpy())

def normalize_text(text):
    return re.sub(r'\s+', ' ', text.strip().lower())

def highlight_relevant_text(query, text, doc):
    QA_input = {'question': query, 'context': text}
    res = nlp(QA_input)
    para = doc.add_paragraph()
    matches, covered_idx = [], set()
    for answer in res:
        for match in re.finditer(re.escape(answer['answer']), text):
            start_idx, end_idx = match.start(), match.end()
            if not any(idx in covered_idx for idx in range(start_idx, end_idx)):
                matches.append((start_idx, end_idx))
                covered_idx.update(range(start_idx, end_idx))
    matches.sort()
    last_index = 0
    for start_idx, end_idx in matches:
        if last_index < start_idx:
            para.add_run(text[last_index:start_idx])
        run = para.add_run(text[start_idx:end_idx])
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        last_index = end_idx
    if last_index < len(text):
        para.add_run(text[last_index:])

def merge_snippets_based_on_content(snippets, MIN_OVERLAP):
    if not snippets:
        return []
    merged = [snippets[0]]
    for current in snippets[1:]:
        previous = merged[-1]
        prev_words, current_words = previous.split(), current.split()
        overlap = 0
        for i in range(1, min(len(prev_words), len(current_words)) + 1):
            if prev_words[-i:] == current_words[:i]:
                overlap = i
        if overlap >= MIN_OVERLAP:
            merged[-1] = ' '.join(prev_words + current_words[overlap:])
        else:
            merged.append(current)
    return merged

# Process files
input_dir = "texts"
text_files = [f for f in os.listdir(input_dir) if f.endswith(".txt")]
batch_results = []
query_embeddings = encode_texts(queries)

for filename in tqdm(text_files, desc="Processing text files"):
    with open(os.path.join(input_dir, filename), 'r', encoding='utf-8') as f:
        full_text = f.read()
    snippets = [full_text[i:i+WINDOW_SIZE] for i in range(0, max(1, len(full_text) - WINDOW_SIZE + 1), WINDOW_SIZE - OVERLAP)]
    merged_snippets = merge_snippets_based_on_content(snippets, MIN_OVERLAP)
    seen = set()
    for snippet in merged_snippets:
        norm_snippet = normalize_text(snippet)
        if norm_snippet in seen:
            continue
        seen.add(norm_snippet)
        snippet_embedding = encode_texts([snippet])
        similarities = calculate_similarity(query_embeddings, snippet_embedding)
        for k in range(len(queries)):
            similarity = similarities[k, 0]
            row = [queries[k], similarity, filename, snippet] + list(query_df.iloc[k, 1:])
            batch_results.append(row)

columns = ["Query", "Similarity", "Filename", "Extract"] + category_columns
df = pd.DataFrame(batch_results, columns=columns)
df_sorted = df.sort_values(by=["Query", "Similarity"], ascending=[True, False])
df_top_emb = df_sorted.groupby("Query").head(TOP_EMB_LIMIT)
df_top_emb["Normalized_Extract"] = df_top_emb["Extract"].apply(normalize_text)
df_top_emb = df_top_emb.loc[df_top_emb.groupby(["Filename", "Normalized_Extract"])["Similarity"].idxmax()].drop(columns=["Normalized_Extract"])

# Save to Excel
output_file = "output_results.xlsx"
with pd.ExcelWriter(output_file) as writer:
    df_top_emb.to_excel(writer, sheet_name="Top Matches", index=False)

# Generate Word reports
def process_and_generate_word(input_file, output_dir):
    df = pd.read_excel(input_file)
    grouped = df.groupby("Query")
    for query, group in grouped:
        parts = [str(p) for p in group.iloc[0][category_columns].tolist() if pd.notna(p)]
        if not parts:
            parts = ["Uncategorized"]
        query_dir = os.path.join(output_dir, *parts)
        os.makedirs(query_dir, exist_ok=True)
        safe_query = query.translate(str.maketrans('', '', string.punctuation)).replace(' ', '_')
        doc = Document()
        doc.add_heading(query, level=1)
        for filename, file_group in group.groupby("Filename"):
            doc.add_heading(filename, level=2)
            for _, row in file_group.iterrows():
                highlight_relevant_text(query, row["Extract"], doc)
                doc.add_paragraph("\n---\n")
        doc.save(os.path.join(query_dir, f"{safe_query}.docx"))
        print(f"Saved: {os.path.join(query_dir, f'{safe_query}.docx')}")

process_and_generate_word(output_file, "analysis_word")
